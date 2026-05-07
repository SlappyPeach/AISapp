from __future__ import annotations

import mimetypes
import re
from datetime import datetime
from decimal import Decimal
from pathlib import Path
from typing import Any, Iterable
from zipfile import ZIP_DEFLATED, ZipFile
from xml.sax.saxutils import escape as xml_escape

from django.conf import settings

from .models import (
    DocumentRecord,
    PPEIssuance,
    PrimaryDocument,
    ProcurementRequest,
    SiteMaterialRequest,
    SMRContract,
    StockIssue,
    StockReceipt,
    SupplierDocument,
    SupplyContract,
    WorkAcceptanceAct,
    WriteOffAct,
)
from .reporting import REPORT_PROVIDERS


def money(value: Any) -> str:
    numeric = float(value or 0)
    return f"{numeric:,.2f}".replace(",", " ").replace(".", ",")


PLACEHOLDER_RE = re.compile(r"\{\{\s*([A-Z0-9_]+)\s*\}\}")
MONTH_NAMES = [
    "января",
    "февраля",
    "марта",
    "апреля",
    "мая",
    "июня",
    "июля",
    "августа",
    "сентября",
    "октября",
    "ноября",
    "декабря",
]

DOCX_TEMPLATE_FILES = {
    "smr_contract": "Договор на СМР_шаблон.docx",
    "supply_contract": "Договор поставки_шаблон.docx",
    "stock_receipt": "Приходный ордер_шаблон.docx",
    "stock_issue": "Требование-накладная_шаблон.docx",
    "write_off": "Акт списания материалов по договору_шаблон.docx",
    "work_acceptance": "Акт сдачи-приемки выполненных работ_шаблон.docx",
}

PRIMARY_DOCUMENT_TEMPLATE_FILES = {
    "invoice": "Счет на материал_шаблон.docx",
    "invoice_facture": "Счет-фактура_шаблон.docx",
    "vat_invoice": "Счет-фактура_шаблон.docx",
    "goods_waybill": "Товарная накладная ТОРГ-12_шаблон.docx",
    "upd": "Товарная накладная ТОРГ-12_шаблон.docx",
    "receipt_invoice": "Товарная накладная ТОРГ-12_шаблон.docx",
}

XLSX_TEMPLATE_FILES = {
    "site_material_report": "Материальный отчет_шаблон.xlsx",
}


def _load_xlsxwriter():
    try:
        import xlsxwriter
    except ImportError as exc:
        raise RuntimeError("Экспорт XLSX недоступен: пакет xlsxwriter не установлен или поврежден.") from exc
    return xlsxwriter


def _load_docx_dependencies():
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError("Экспорт DOCX недоступен: пакет python-docx или lxml не установлен либо поврежден.") from exc
    return Document, WD_ALIGN_PARAGRAPH, Pt


class Exporter:
    def content_type(self, path: Path) -> str:
        guessed, _ = mimetypes.guess_type(path.name)
        return guessed or "application/octet-stream"

    def _template_path(self, template_name: str) -> Path | None:
        templates_dir = Path(getattr(settings, "DOCUMENT_TEMPLATES_DIR", ""))
        path = templates_dir / template_name
        return path if path.exists() else None

    def _render_docx_template(self, template_name: str, context: dict[str, Any], path: Path) -> bool:
        template_path = self._template_path(template_name)
        if not template_path:
            return False

        replacements = {key: "" if value is None else str(value) for key, value in context.items()}

        def replace_placeholders(raw_xml: bytes) -> bytes:
            xml_text = raw_xml.decode("utf-8")

            def replace(match: re.Match[str]) -> str:
                return xml_escape(replacements.get(match.group(1), ""))

            return PLACEHOLDER_RE.sub(replace, xml_text).encode("utf-8")

        with ZipFile(template_path, "r") as source, ZipFile(path, "w", ZIP_DEFLATED) as target:
            for member in source.infolist():
                data = source.read(member.filename)
                if member.filename.startswith("word/") and member.filename.endswith(".xml"):
                    data = replace_placeholders(data)
                target.writestr(member, data)
        return True

    def _render_xlsx_template(self, template_name: str, context: dict[str, Any], path: Path) -> bool:
        template_path = self._template_path(template_name)
        if not template_path:
            return False

        replacements = {key: "" if value is None else str(value) for key, value in context.items()}

        def replace_placeholders(raw_xml: bytes) -> bytes:
            xml_text = raw_xml.decode("utf-8")

            def replace(match: re.Match[str]) -> str:
                return xml_escape(replacements.get(match.group(1), ""))

            return PLACEHOLDER_RE.sub(replace, xml_text).encode("utf-8")

        with ZipFile(template_path, "r") as source, ZipFile(path, "w", ZIP_DEFLATED) as target:
            for member in source.infolist():
                data = source.read(member.filename)
                if member.filename.endswith(".xml"):
                    data = replace_placeholders(data)
                target.writestr(member, data)
        return True

    def _date_text(self, value: Any) -> str:
        if not value:
            return ""
        return value.strftime("%d.%m.%Y") if hasattr(value, "strftime") else str(value)

    def _date_parts(self, prefix: str, value: Any) -> dict[str, str]:
        if not value or not hasattr(value, "month"):
            return {f"{prefix}_DAY": "", f"{prefix}_MONTH": "", f"{prefix}_YEAR": ""}
        return {
            f"{prefix}_DAY": f"{value.day:02d}",
            f"{prefix}_MONTH": MONTH_NAMES[value.month - 1],
            f"{prefix}_YEAR": str(value.year),
        }

    def _duration_days(self, start_date: Any, end_date: Any) -> str:
        if not start_date or not end_date:
            return ""
        try:
            return str((end_date - start_date).days + 1)
        except TypeError:
            return ""

    def _template_common_context(self) -> dict[str, str]:
        profile = self._organization_profile()
        organization_name = profile["name"] or "АО «СТ-1»"
        return {
            "ORGANIZATION_NAME": organization_name,
            "CITY": "Москва",
            "CONTRACT_CITY": "Москва",
            "DIRECTOR_NAME": "________________",
            "RESPONSIBLE_PERSON_NAME": "________________",
            "SITE_MANAGER_NAME": "________________",
            "LEFT_SIGNER_NAME": "________________",
            "RIGHT_SIGNER_NAME": "________________",
            "CONTRACTOR_SIGNER_NAME": "________________",
            "CONTRACTOR_SIGNER_POSITION": "представителя",
            "CONTRACTOR_AUTH_DOC": "доверенности",
            "CUSTOMER_SIGNER_NAME": "________________",
            "CUSTOMER_SIGNER_POSITION": "представителя",
            "CUSTOMER_AUTH_DOC": "доверенности",
            "CUSTOMER_SIGNER_AUTH_DOC": "доверенности",
            "BUYER_SIGNER": "________________",
            "BUYER_AUTH_DOC": "доверенности",
            "SUPPLIER_SIGNER": "________________",
            "SUPPLIER_AUTH_DOC": "доверенности",
            "JURISDICTION_PARTY": organization_name,
        }

    def _add_line_context(self, context: dict[str, Any], lines: Iterable[Any], mapper, *, limit: int = 12) -> None:
        for index, line in enumerate(list(lines)[:limit], start=1):
            suffix = "" if index == 1 else f"_{index}"
            for key, value in mapper(line, index).items():
                context[f"{key}{suffix}"] = value

    def _line_amount(self, line: Any) -> Decimal:
        return Decimal(line.quantity or 0) * Decimal(line.unit_price or 0)

    def _xlsx_path(self, prefix: str, date_from: Any, date_to: Any) -> Path:
        return settings.EXPORTS_DIR / f"{prefix}_{date_from}_{date_to}.xlsx"

    def _site_material_report_template_context(self, rows: list[dict[str, Any]], filters: dict[str, Any], *, user=None) -> dict[str, Any]:
        date_from = filters.get("date_from") or datetime.now().date().replace(day=1)
        date_to = filters.get("date_to") or datetime.now().date()
        detail_rows = [row for row in rows if str(row.get("Период", "")).upper() != "ИТОГО"]
        site_name = (
            (getattr(user, "site_name", "") or "").strip()
            or (filters.get("location_name") or "").strip()
            or (detail_rows[0].get("Место хранения", "") if detail_rows else "")
        )

        context: dict[str, Any] = {
            **self._template_common_context(),
            "REPORT_NUMBER": f"MAT-{date_from:%Y%m}",
            "REPORT_PERIOD": f"{self._date_text(date_from)} - {self._date_text(date_to)}",
            "SITE_NAME": site_name or "-",
            "RESPONSIBLE_PERSON_NAME": getattr(user, "full_name_or_username", "") if user is not None else "",
        }

        for index, row in enumerate(detail_rows[:12], start=1):
            suffix = "" if index == 1 else f"_{index}"
            context.update(
                {
                    f"LINE_NO{suffix}": index,
                    f"MATERIAL_CODE{suffix}": row.get("Код материала", ""),
                    f"MATERIAL_NAME{suffix}": row.get("Наименование материала", ""),
                    f"UNIT{suffix}": row.get("Ед. изм.", ""),
                    f"OPENING_QTY{suffix}": row.get("Остаток на начало", ""),
                    f"RECEIPT_QTY{suffix}": row.get("Поступило за период", ""),
                    f"ISSUE_QTY{suffix}": row.get("Израсходовано за период", ""),
                    f"CLOSING_QTY{suffix}": row.get("Остаток на конец", ""),
                    f"PRICE{suffix}": row.get("Цена за единицу", ""),
                    f"CLOSING_AMOUNT{suffix}": row.get("Сумма остатка", ""),
                    f"BASIS_DOCUMENT{suffix}": row.get("Период", ""),
                }
            )
        return context

    def export_document(self, entity_type: str, entity_id: int) -> Path:
        handlers = {
            "smr_contract": self._export_smr_contract,
            "supply_contract": self._export_supply_contract,
            "site_material_request": self._export_site_material_request,
            "procurement_request": self._export_procurement_request,
            "primary_document": self._export_primary_document,
            "stock_receipt": self._export_stock_receipt,
            "stock_issue": self._export_stock_issue,
            "write_off": self._export_writeoff,
            "ppe_issuance": self._export_ppe_issuance,
            "supplier_document": self._export_supplier_document,
            "work_acceptance": self._export_work_acceptance,
        }
        if entity_type not in handlers:
            raise ValueError("Для этого типа документа выгрузка не реализована.")
        path = handlers[entity_type](entity_id)
        DocumentRecord.objects.filter(entity_type=entity_type, entity_id=entity_id).update(file_path=str(path))
        return path

    def export_report(self, report_name: str, filters: dict[str, Any], *, user=None) -> Path:
        provider = REPORT_PROVIDERS[report_name]
        rows = provider(filters, user=user)
        date_from = filters.get("date_from") or datetime.now().date().replace(day=1)
        date_to = filters.get("date_to") or datetime.now().date()
        path = self._xlsx_path(report_name, date_from, date_to)

        template_name = XLSX_TEMPLATE_FILES.get(report_name)
        if template_name and self._render_xlsx_template(
            template_name,
            self._site_material_report_template_context(rows, filters, user=user),
            path,
        ):
            return path

        xlsxwriter = _load_xlsxwriter()
        workbook = xlsxwriter.Workbook(str(path))
        worksheet = workbook.add_worksheet("Отчет")
        title_format = workbook.add_format({"bold": True, "font_size": 14})
        header_format = workbook.add_format({"bold": True, "bg_color": "#E6D7C6", "border": 1})
        cell_format = workbook.add_format({"border": 1})
        numeric_format = workbook.add_format({"border": 1, "num_format": "#,##0.00"})
        total_cell_format = workbook.add_format({"border": 1, "bold": True, "bg_color": "#F3E7D7"})
        total_numeric_format = workbook.add_format({"border": 1, "bold": True, "bg_color": "#F3E7D7", "num_format": "#,##0.00"})

        worksheet.write("A1", "Экспорт отчета АИС", title_format)
        worksheet.write("A2", f"Период: {date_from} - {date_to}")
        worksheet.freeze_panes(4, 0)

        if rows:
            headers = list(rows[0].keys())
            for col_index, header in enumerate(headers):
                worksheet.write(3, col_index, header, header_format)
            for row_index, row in enumerate(rows, start=4):
                first_header = headers[0]
                first_value = row.get(first_header, "")
                is_total_row = str(first_value).startswith("ИТОГО")
                for col_index, header in enumerate(headers):
                    value = row.get(header)
                    if isinstance(value, Decimal):
                        value = float(value)
                    if isinstance(value, (int, float)):
                        worksheet.write_number(
                            row_index,
                            col_index,
                            float(value),
                            total_numeric_format if is_total_row else numeric_format,
                        )
                    else:
                        worksheet.write(
                            row_index,
                            col_index,
                            value,
                            total_cell_format if is_total_row else cell_format,
                        )
            for col_index, header in enumerate(headers):
                max_len = max(len(str(header)), *(len(str(item.get(header, ""))) for item in rows))
                worksheet.set_column(col_index, col_index, min(max_len + 2, 45))
            worksheet.autofilter(3, 0, 3 + len(rows), len(headers) - 1)
        else:
            worksheet.write("A4", "Нет данных за выбранный период.")

        workbook.close()
        return path

    def _doc_path(self, prefix: str, number: str) -> Path:
        safe_number = number.replace("/", "_").replace("\\", "_").replace(" ", "_")
        return settings.EXPORTS_DIR / f"{prefix}_{safe_number}.docx"

    def _prepare_doc(self, title: str, subtitle: str = ""):
        Document, WD_ALIGN_PARAGRAPH, Pt = _load_docx_dependencies()
        document = Document()
        style = document.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)

        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(14)

        if subtitle:
            subtitle_paragraph = document.add_paragraph()
            subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_paragraph.add_run(subtitle).italic = True

        document.add_paragraph()
        return document

    def _add_meta(self, document, items: Iterable[tuple[str, str]]) -> None:
        for label, value in items:
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{label}: ").bold = True
            paragraph.add_run(value)

    def _add_heading(self, document, text: str) -> None:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = _load_docx_dependencies()[2](12)

    def _add_clause(self, document, number: str, text: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{number}. ").bold = True
        paragraph.add_run(text)

    def _add_table(self, document, headers: list[str], rows: list[list[str]]) -> None:
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = value

    def _add_signature(self, document, left_label: str, right_label: str) -> None:
        document.add_paragraph()
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.cell(0, 0).text = f"{left_label}\n\n_____________________"
        table.cell(0, 1).text = f"{right_label}\n\n_____________________"

    def _organization_profile(self) -> dict[str, str]:
        profile = getattr(settings, "ORGANIZATION_PROFILE", {}) or {}
        if not isinstance(profile, dict):
            profile = {"name": str(profile)}
        return {
            "name": str(profile.get("name", "")).strip(),
            "tax_id": str(profile.get("tax_id", "")).strip(),
            "kpp": str(profile.get("kpp", "")).strip(),
            "ogrn": str(profile.get("ogrn", "")).strip(),
            "address": str(profile.get("address", "")).strip(),
            "bank_details": str(profile.get("bank_details", "")).strip(),
            "requisites": str(profile.get("requisites", "")).strip(),
        }

    def _organization_name(self) -> str:
        profile = self._organization_profile()
        return profile["name"]

    def _organization_requisites(self) -> str:
        profile = self._organization_profile()
        if profile["requisites"]:
            return profile["requisites"]

        parts: list[str] = []
        if profile["tax_id"]:
            parts.append(f"ИНН {profile['tax_id']}")
        if profile["kpp"]:
            parts.append(f"КПП {profile['kpp']}")
        if profile["ogrn"]:
            parts.append(f"ОГРН {profile['ogrn']}")
        if profile["address"]:
            parts.append(profile["address"])
        if profile["bank_details"]:
            parts.append(profile["bank_details"])
        return "; ".join(part for part in parts if part)

    def _supplier_requisites(self, supplier) -> str:
        if hasattr(supplier, "requisites_text"):
            return supplier.requisites_text()
        return ""

    def _smr_contract_template_context(self, contract: SMRContract) -> dict[str, Any]:
        object_name = contract.object.name if contract.object else ""
        vat_amount = Decimal(contract.amount or 0) * Decimal(contract.vat_rate or 0) / Decimal("100")
        context = {
            **self._template_common_context(),
            **self._date_parts("CONTRACT", contract.contract_date),
            "CONTRACT_NUMBER": contract.number,
            "CUSTOMER_NAME": contract.resolved_customer_name() or "-",
            "CUSTOMER_REQUISITES": contract.resolved_customer_requisites() or "-",
            "CONTRACTOR_NAME": contract.resolved_contractor_name() or "-",
            "CONTRACTOR_REQUISITES": contract.resolved_contractor_requisites() or "-",
            "OBJECT_NAME": object_name,
            "OBJECT_ADDRESS": contract.object.address if contract.object else "",
            "WORK_SUBJECT": contract.subject,
            "WORK_LINE_NO": "1",
            "WORK_NAME": contract.work_type or contract.subject,
            "WORK_QUANTITY": contract.planned_volume or "",
            "WORK_UNIT": contract.volume_unit or "",
            "ESTIMATE_LINE": f"{contract.work_type or contract.subject}: {contract.planned_volume or '-'} {contract.volume_unit or ''}".strip(),
            "CONTRACT_AMOUNT_TEXT": f"{money(contract.amount)} руб.",
            "VAT_AMOUNT_TEXT": f"{money(vat_amount)} руб.",
            "WORK_DURATION_DAYS": self._duration_days(contract.start_date, contract.end_date),
        }
        return context

    def _supply_contract_template_context(self, contract: SupplyContract) -> dict[str, Any]:
        vat_rate = Decimal("20")
        vat_amount = Decimal(contract.amount or 0) * vat_rate / Decimal("100")
        buyer_name = self._organization_name() or "-"
        return {
            **self._template_common_context(),
            **self._date_parts("CONTRACT", contract.contract_date),
            "CONTRACT_NUMBER": contract.number,
            "SUPPLIER_NAME": contract.supplier.name,
            "SUPPLIER_REQUISITES": self._supplier_requisites(contract.supplier) or "-",
            "BUYER_NAME": buyer_name,
            "BUYER_REQUISITES": self._organization_requisites() or "-",
            "CONTRACT_AMOUNT": money(contract.amount),
            "VAT_RATE": str(vat_rate),
            "VAT_AMOUNT": money(vat_amount),
            "DELIVERY_PLACE": settings.WAREHOUSE_NAME,
            "DELIVERY_TERM": contract.terms or "по согласованным заявкам",
            "PAYMENT_TERM": "по условиям договора",
            "VALID_UNTIL": "",
            "LEFT_SIGNER_NAME": contract.supplier.name,
            "RIGHT_SIGNER_NAME": buyer_name,
        }

    def _primary_document_template_context(self, item: PrimaryDocument) -> dict[str, Any]:
        lines = list(item.lines.all())
        total_amount = Decimal(item.amount or 0)
        if not total_amount:
            total_amount = sum((self._line_amount(line) for line in lines), Decimal("0"))
        vat_amount = Decimal(item.vat_amount or 0)
        supplier_requisites = self._supplier_requisites(item.supplier) or "-"
        buyer_name = self._organization_name() or "-"
        buyer_requisites = self._organization_requisites() or "-"
        supplier_inn = item.supplier.tax_id or ""

        context: dict[str, Any] = {
            **self._template_common_context(),
            "INVOICE_NUMBER": item.number,
            "INVOICE_DATE": self._date_text(item.doc_date),
            "INVOICE_FACTURE_NUMBER": item.number,
            "INVOICE_FACTURE_DATE": self._date_text(item.doc_date),
            "WAYBILL_NUMBER": item.number,
            "WAYBILL_DATE": self._date_text(item.doc_date),
            "BASIS_DOCUMENT": item.basis_reference or (item.procurement_request.number if item.procurement_request else ""),
            "PAYMENT_DOCUMENT": item.basis_reference or "",
            "SUPPLIER_NAME": item.supplier.name,
            "SELLER_NAME": item.supplier.name,
            "SHIPPER_NAME": item.supplier.name,
            "SUPPLIER_REQUISITES": supplier_requisites,
            "SELLER_ADDRESS": item.supplier.address or supplier_requisites,
            "SHIPPER_REQUISITES": supplier_requisites,
            "SUPPLIER_INN": supplier_inn,
            "SUPPLIER_KPP": "",
            "SELLER_INN_KPP": supplier_inn,
            "BUYER_NAME": buyer_name,
            "BUYER_REQUISITES": buyer_requisites,
            "BUYER_ADDRESS": self._organization_profile().get("address", ""),
            "BUYER_INN_KPP": self._organization_profile().get("tax_id", ""),
            "CONSIGNEE_NAME": buyer_name,
            "CONSIGNEE_NAME_ADDRESS": f"{buyer_name}, {self._organization_profile().get('address', '')}".strip(", "),
            "CONSIGNEE_REQUISITES": buyer_requisites,
            "PAYER_NAME": buyer_name,
            "PAYER_REQUISITES": buyer_requisites,
            "ITEMS_COUNT": len(lines),
            "TOTAL_AMOUNT": money(total_amount),
            "TOTAL_TO_PAY": money(total_amount),
            "TOTAL_TO_PAY_WORDS": f"{money(total_amount)} руб.",
            "VAT_RATE": "20",
            "VAT_AMOUNT": money(vat_amount),
            "AMOUNT_NO_VAT": money(total_amount - vat_amount),
            "AMOUNT_WITH_VAT": money(total_amount),
            "CURRENCY_NAME": "Российский рубль",
            "CURRENCY_CODE": "643",
            "COUNTRY": "Россия",
            "EXCISE": "без акциза",
            "CUSTOMS_DECLARATION": "-",
            "PACKING_TYPE": "-",
            "PACKS_COUNT": "",
            "WEIGHT": "",
            "COMMENT": item.notes or "",
            "RECEIVER_BANK": "",
            "BANK_BIK": "",
            "BANK_ACCOUNT": "",
        }

        def map_line(line, index: int) -> dict[str, Any]:
            line_amount = self._line_amount(line)
            return {
                "LINE_NO": index,
                "MATERIAL_CODE": line.material.code,
                "ITEM_CODE": line.material.code,
                "ITEM_NAME": line.material.name,
                "UNIT": line.material.unit,
                "QUANTITY": line.quantity,
                "PRICE": money(line.unit_price),
                "LINE_AMOUNT": money(line_amount),
                "AMOUNT_NO_VAT": money(line_amount),
                "AMOUNT_WITH_VAT": money(line_amount),
                "VAT_AMOUNT": "",
            }

        self._add_line_context(context, lines, map_line)
        return context

    def _stock_receipt_template_context(self, receipt: StockReceipt) -> dict[str, Any]:
        lines = list(receipt.lines.all())
        total_amount = sum((self._line_amount(line) for line in lines), Decimal("0"))
        total_qty = sum((Decimal(line.quantity or 0) for line in lines), Decimal("0"))
        supplier_document = receipt.supplier_document
        primary_document = receipt.primary_document
        context: dict[str, Any] = {
            **self._template_common_context(),
            "RECEIPT_ORDER_NUMBER": receipt.number,
            "DOCUMENT_DATE": self._date_text(receipt.receipt_date),
            "ORGANIZATION_NAME": self._organization_name() or "АО «СТ-1»",
            "STRUCTURAL_UNIT": settings.WAREHOUSE_NAME,
            "WAREHOUSE": settings.WAREHOUSE_NAME,
            "SUPPLIER_NAME": receipt.supplier.name,
            "SUPPLIER_DOCUMENT_NUMBER": supplier_document.doc_number if supplier_document else (primary_document.number if primary_document else ""),
            "SUPPLIER_DOCUMENT_DATE": self._date_text(supplier_document.doc_date if supplier_document else (primary_document.doc_date if primary_document else "")),
            "PAYMENT_DOCUMENT_NUMBER": primary_document.number if primary_document else "",
            "TOTAL_DOCUMENT_QTY": total_qty,
            "TOTAL_ACCEPTED_QTY": total_qty,
            "TOTAL_AMOUNT_NO_VAT": money(total_amount),
            "TOTAL_VAT": "",
            "TOTAL_WITH_VAT": money(total_amount),
            "OKPO": "",
            "ACCOUNT_CODE": "",
            "STOCK_CARD_NUMBER": "",
            "SENDER_POSITION": "Поставщик",
            "RECEIVER_POSITION": "Кладовщик",
        }

        def map_line(line, index: int) -> dict[str, Any]:
            return {
                "MATERIAL_CODE": line.material.code,
                "MATERIAL_NAME": line.material.name,
                "UNIT": line.material.unit,
                "DOCUMENT_QTY": line.quantity,
                "ACCEPTED_QTY": line.quantity,
                "PRICE": money(line.unit_price),
                "AMOUNT_NO_VAT": money(self._line_amount(line)),
                "VAT_AMOUNT": "",
            }

        self._add_line_context(context, lines, map_line)
        return context

    def _stock_issue_template_context(self, issue: StockIssue) -> dict[str, Any]:
        lines = list(issue.lines.all())
        context: dict[str, Any] = {
            **self._template_common_context(),
            "DOCUMENT_NUMBER": issue.number,
            "DOCUMENT_DATE": self._date_text(issue.issue_date),
            "SENDER_UNIT": settings.WAREHOUSE_NAME,
            "RECEIVER_UNIT": issue.site_name,
            "SENDER_ACTIVITY": "склад",
            "RECEIVER_ACTIVITY": "строительно-монтажные работы",
            "RECEIVED_BY_POSITION": issue.received_by_name,
            "ISSUED_BY_POSITION": "Кладовщик",
            "OPERATION_CODE": "",
            "ACCOUNTING_UNIT": "",
            "ACCOUNT_CODE": "",
        }

        def map_line(line, index: int) -> dict[str, Any]:
            return {
                "MATERIAL_CODE": line.material.code,
                "MATERIAL_NAME": line.material.name,
                "UNIT": line.material.unit,
                "REQUESTED_QTY": line.quantity,
                "ISSUED_QTY": line.quantity,
                "PRICE": money(line.unit_price),
                "AMOUNT": money(self._line_amount(line)),
                "STOCK_CARD_NUMBER": "",
            }

        self._add_line_context(context, lines, map_line)
        return context

    def _writeoff_template_context(self, act: WriteOffAct) -> dict[str, Any]:
        context: dict[str, Any] = {
            **self._template_common_context(),
            **self._date_parts("APPROVAL", act.act_date),
            "ACT_NUMBER": act.number,
            "MONTH": MONTH_NAMES[act.act_date.month - 1],
            "YEAR": act.act_date.year,
            "SITE_NUMBER": act.site_name,
            "CONTRACT_NUMBER": act.contract.number,
            "CONTRACT_DATE": self._date_text(act.contract.contract_date),
            "CONTRACT_SUBJECT": act.contract.subject,
            "CONTRACT_WORK_COLUMN": act.work_type,
        }

        def map_line(line, index: int) -> dict[str, Any]:
            return {
                "LINE_NO": index,
                "MATERIAL_CODE": line.material.code,
                "MATERIAL_NAME": line.material.name,
                "REPORT_MATERIAL_NAME": line.material.name,
                "UNIT": line.material.unit,
                "NORM_QTY": line.calculated_quantity,
                "ACTUAL_QTY": line.actual_quantity,
                "WORK_OR_FORM": act.work_type,
            }

        self._add_line_context(context, list(act.lines.all()), map_line)
        return context

    def _ppe_template_context(self, issuance: PPEIssuance) -> dict[str, Any]:
        context: dict[str, Any] = {
            **self._template_common_context(),
            "DOCUMENT_NUMBER": issuance.number,
            "DOCUMENT_DATE": self._date_text(issuance.issue_date),
            "SITE_NAME": issuance.site_name,
            "PPE_CATEGORY": issuance.season or "СИЗ",
        }

        def map_line(line, index: int) -> dict[str, Any]:
            sizes = []
            if line.clothing_size:
                sizes.append(f"размер одежды {line.clothing_size}")
            if line.shoe_size:
                sizes.append(f"размер обуви {line.shoe_size}")
            size_note = f" ({', '.join(sizes)})" if sizes else ""
            return {
                "LINE_NO": index,
                "WORKER_NAME": line.worker.full_name,
                "EMPLOYEE_NUMBER": line.worker.employee_number,
                "PPE_NAME": f"{line.material.name}{size_note}",
                "PPE_CODE": line.material.code,
                "NOMENCLATURE_NUMBER": line.material.code,
                "UNIT": line.material.unit,
                "UNIT_NAME": line.material.unit,
                "UNIT_CODE": line.material.unit,
                "QUANTITY": line.quantity,
                "SERVICE_LIFE_MONTHS": line.service_life_months,
                "START_DATE": self._date_text(line.replacement_start_date),
                "WORKER_SIGNATURE": "",
            }

        self._add_line_context(context, list(issuance.lines.all()), map_line, limit=16)
        return context

    def _work_acceptance_template_context(self, act: WorkAcceptanceAct) -> dict[str, Any]:
        vat_rate = Decimal(act.contract.vat_rate or 0)
        vat_amount = Decimal(act.amount or 0) * vat_rate / Decimal("100")
        return {
            **self._template_common_context(),
            **self._date_parts("ACT", act.act_date),
            **self._date_parts("CONTRACT", act.contract.contract_date),
            "ACT_NUMBER": act.number,
            "CONTRACT_NUMBER": act.contract.number,
            "CUSTOMER_NAME": act.contract.resolved_customer_name() or "-",
            "CONTRACTOR_NAME": act.contract.resolved_contractor_name() or self._organization_name() or "-",
            "OBJECT_ADDRESS": act.contract.object.address if act.contract.object else act.site_name,
            "WORK_DESCRIPTION": act.work_description or act.contract.subject,
            "WORK_PERIOD": f"{self._date_text(act.contract.start_date)} - {self._date_text(act.contract.end_date)}",
            "AMOUNT": money(act.amount),
            "VAT_RATE": vat_rate,
            "VAT_AMOUNT": money(vat_amount),
            "COPIES_COUNT": "2",
            "LEFT_SIGNER_NAME": act.contract.resolved_customer_name() or "-",
            "RIGHT_SIGNER_NAME": act.contract.resolved_contractor_name() or self._organization_name() or "-",
        }

    def _supplier_document_template_name(self, doc_type: str) -> str | None:
        normalized = doc_type.casefold()
        if "счет-фактура" in normalized or "счёт-фактура" in normalized:
            return "Счет-фактура_шаблон.docx"
        if "наклад" in normalized:
            return "Товарная накладная ТОРГ-12_шаблон.docx"
        if "счет" in normalized or "счёт" in normalized:
            return "Счет на оплату по скану_шаблон.docx"
        return None

    def _export_smr_contract(self, entity_id: int) -> Path:
        contract = SMRContract.objects.select_related("object").get(pk=entity_id)
        path = self._doc_path("smr_contract", contract.number)
        if self._render_docx_template(DOCX_TEMPLATE_FILES["smr_contract"], self._smr_contract_template_context(contract), path):
            return path
        customer_name = contract.resolved_customer_name() or "-"
        customer_requisites = contract.resolved_customer_requisites() or "-"
        contractor_name = contract.resolved_contractor_name() or "-"
        contractor_requisites = contract.resolved_contractor_requisites() or "-"
        doc = self._prepare_doc("ДОГОВОР НА ВЫПОЛНЕНИЕ СМР", f"№ {contract.number} от {contract.contract_date}")
        self._add_meta(
            doc,
            [
                ("Заказчик", customer_name),
                ("Реквизиты заказчика", customer_requisites),
                ("Подрядчик", contractor_name),
                ("Реквизиты подрядчика", contractor_requisites),
                ("Объект", contract.object.name if contract.object else ""),
                ("Предмет", contract.subject),
                ("Вид работ", contract.work_type or ""),
                ("Плановый объем", f"{contract.planned_volume or 0} {contract.volume_unit or ''}".strip()),
                ("Стоимость", f"{money(contract.amount)} руб."),
                ("Сроки", f"{contract.start_date or '-'} - {contract.end_date or '-'}"),
            ],
        )
        self._add_heading(doc, "1. Предмет договора")
        self._add_clause(
            doc,
            "1.1",
            f"Подрядчик обязуется выполнить работы «{contract.subject}» на объекте "
            f"«{contract.object.name if contract.object else '-'}», а Заказчик обязуется принять и оплатить результат работ.",
        )
        self._add_clause(
            doc,
            "1.2",
            f"Вид работ: {contract.work_type or '-'}. Плановый объем: "
            f"{contract.planned_volume or '-'} {contract.volume_unit or ''}.",
        )
        self._add_heading(doc, "2. Стоимость и сроки")
        self._add_clause(doc, "2.1", f"Стоимость работ составляет {money(contract.amount)} руб., НДС {contract.vat_rate}%.")
        self._add_clause(doc, "2.2", f"Срок выполнения работ: с {contract.start_date or '-'} по {contract.end_date or '-'}.")
        self._add_signature(doc, f"Заказчик: {customer_name}", f"Подрядчик: {contractor_name}")
        doc.save(path)
        return path

    def _export_supply_contract(self, entity_id: int) -> Path:
        contract = SupplyContract.objects.select_related("supplier", "related_smr_contract").get(pk=entity_id)
        path = self._doc_path("supply_contract", contract.number)
        if self._render_docx_template(DOCX_TEMPLATE_FILES["supply_contract"], self._supply_contract_template_context(contract), path):
            return path
        buyer_name = self._organization_name() or "-"
        buyer_requisites = self._organization_requisites() or "-"
        supplier_requisites = self._supplier_requisites(contract.supplier) or "-"
        doc = self._prepare_doc("ДОГОВОР ПОСТАВКИ", f"№ {contract.number} от {contract.contract_date}")
        self._add_meta(
            doc,
            [
                ("Поставщик", contract.supplier.name),
                ("Реквизиты поставщика", supplier_requisites),
                ("Покупатель", buyer_name),
                ("Реквизиты покупателя", buyer_requisites),
                ("Связанный договор СМР", contract.related_smr_contract.number if contract.related_smr_contract else "-"),
                ("Сумма", f"{money(contract.amount)} руб."),
                ("Статус", contract.get_status_display()),
            ],
        )
        doc.add_paragraph(contract.terms or "Поставка материалов выполняется по заявкам снабженца через АИС.")
        self._add_signature(doc, f"Поставщик: {contract.supplier.name}", f"Покупатель: {buyer_name}")
        doc.save(path)
        return path

    def _export_site_material_request(self, entity_id: int) -> Path:
        request = SiteMaterialRequest.objects.select_related("contract", "requested_by").prefetch_related("lines__material").get(pk=entity_id)
        doc = self._prepare_doc("ЗАЯВКА НА МАТЕРИАЛЫ СО СКЛАДА", f"№ {request.number} от {request.request_date}")
        self._add_meta(
            doc,
            [
                ("Участок", request.site_name),
                ("Договор СМР", request.contract.number if request.contract else "-"),
                ("Заявитель", request.requested_by.full_name_or_username),
                ("Статус", request.get_status_display()),
                ("Комментарий", request.notes or "-"),
            ],
        )
        self._add_table(
            doc,
            ["Код", "Наименование", "Ед.", "Количество", "Цена", "Примечание"],
            [
                [
                    line.material.code,
                    line.material.name,
                    line.material.unit,
                    str(line.quantity),
                    money(line.unit_price),
                    line.notes or "",
                ]
                for line in request.lines.all()
            ],
        )
        self._add_signature(doc, "Начальник участка", "Кладовщик")
        path = self._doc_path("site_material_request", request.number)
        doc.save(path)
        return path

    def _export_procurement_request(self, entity_id: int) -> Path:
        request = ProcurementRequest.objects.select_related("contract", "site_request", "supplier").prefetch_related("lines__material").get(pk=entity_id)
        doc = self._prepare_doc("ЗАЯВКА НА ЗАКУПКУ МАТЕРИАЛОВ", f"№ {request.number} от {request.request_date}")
        self._add_meta(
            doc,
            [
                ("Участок", request.site_name),
                ("Договор СМР", request.contract.number if request.contract else "-"),
                ("Основание", f"Заявка участка {request.site_request.number}" if request.site_request else "-"),
                ("Поставщик", request.supplier.name if request.supplier else "-"),
                ("Статус", request.get_status_display()),
            ],
        )
        self._add_table(
            doc,
            ["Код", "Наименование", "Ед.", "Количество", "Цена", "Примечание"],
            [
                [
                    line.material.code,
                    line.material.name,
                    line.material.unit,
                    str(line.quantity),
                    money(line.unit_price),
                    line.notes or "",
                ]
                for line in request.lines.all()
            ],
        )
        self._add_signature(doc, "Начальник участка / снабженец", "Начальник монтажного объекта")
        path = self._doc_path("procurement_request", request.number)
        doc.save(path)
        return path

    def _export_primary_document(self, entity_id: int) -> Path:
        item = (
            PrimaryDocument.objects.select_related("document_type", "supplier", "procurement_request", "supply_contract", "stock_receipt")
            .prefetch_related("lines__material")
            .get(pk=entity_id)
        )
        path = self._doc_path(item.document_type.code, item.number)
        template_name = PRIMARY_DOCUMENT_TEMPLATE_FILES.get(item.document_type.code)
        if template_name and self._render_docx_template(template_name, self._primary_document_template_context(item), path):
            return path
        receiver_name = self._organization_name() or "-"
        receiver_requisites = self._organization_requisites() or "-"
        supplier_requisites = self._supplier_requisites(item.supplier) or "-"
        title_map = {
            "invoice": "СЧЕТ НА ОПЛАТУ",
            "invoice_facture": "СЧЕТ-ФАКТУРА",
            "upd": "УПД",
            "vat_invoice": "СЧЕТ-ФАКТУРА",
            "goods_waybill": "ТОВАРНАЯ НАКЛАДНАЯ",
            "receipt_invoice": "ПРИХОДНАЯ НАКЛАДНАЯ",
        }
        doc = self._prepare_doc(title_map.get(item.document_type.code, item.document_type.name.upper()), f"№ {item.number} от {item.doc_date}")
        self._add_meta(
            doc,
            [
                ("Тип документа", item.document_type.name),
                ("Поставщик", item.supplier.name),
                ("Реквизиты поставщика", supplier_requisites),
                ("Получатель", receiver_name),
                ("Реквизиты получателя", receiver_requisites),
                ("Основание", item.basis_reference or "-"),
                ("Участок/склад", item.site_name or "-"),
                ("Сумма", f"{money(item.amount)} руб."),
                ("НДС", f"{money(item.vat_amount)} руб."),
                ("Статус", item.get_status_display()),
            ],
        )
        self._add_table(
            doc,
            ["Код", "Наименование", "Ед.", "Количество", "Цена", "Сумма", "Примечание"],
            [
                [
                    line.material.code,
                    line.material.name,
                    line.material.unit,
                    str(line.quantity),
                    money(line.unit_price),
                    money(line.quantity * line.unit_price),
                    line.notes or "",
                ]
                for line in item.lines.all()
            ],
        )
        self._add_signature(doc, f"Поставщик: {item.supplier.name}", f"Получатель: {receiver_name}")
        doc.save(path)
        return path

    def _export_stock_receipt(self, entity_id: int) -> Path:
        receipt = StockReceipt.objects.select_related("supplier", "supplier_document", "primary_document").prefetch_related("lines__material").get(pk=entity_id)
        path = self._doc_path("stock_receipt", receipt.number)
        if self._render_docx_template(DOCX_TEMPLATE_FILES["stock_receipt"], self._stock_receipt_template_context(receipt), path):
            return path
        doc = self._prepare_doc("ПРИХОДНЫЙ ОРДЕР", f"№ {receipt.number} от {receipt.receipt_date}")
        self._add_meta(
            doc,
            [
                ("Поставщик", receipt.supplier.name),
                ("Документ поставщика", receipt.supplier_document.doc_number if receipt.supplier_document else "-"),
                ("Товарная накладная / УПД", receipt.primary_document.number if receipt.primary_document else "-"),
                ("Склад", settings.WAREHOUSE_NAME),
                ("Статус", receipt.get_status_display()),
            ],
        )
        self._add_table(
            doc,
            ["Код", "Наименование", "Ед.", "Количество", "Цена", "Сумма"],
            [
                [
                    line.material.code,
                    line.material.name,
                    line.material.unit,
                    str(line.quantity),
                    money(line.unit_price),
                    money(line.quantity * line.unit_price),
                ]
                for line in receipt.lines.all()
            ],
        )
        self._add_signature(doc, "Кладовщик", "Материально ответственное лицо")
        doc.save(path)
        return path

    def _export_stock_issue(self, entity_id: int) -> Path:
        issue = StockIssue.objects.select_related("contract", "site_request", "stock_receipt").prefetch_related("lines__material").get(pk=entity_id)
        path = self._doc_path("stock_issue", issue.number)
        if self._render_docx_template(DOCX_TEMPLATE_FILES["stock_issue"], self._stock_issue_template_context(issue), path):
            return path
        doc = self._prepare_doc("ТРЕБОВАНИЕ-НАКЛАДНАЯ", f"№ {issue.number} от {issue.issue_date}")
        self._add_meta(
            doc,
            [
                ("Участок", issue.site_name),
                ("Договор СМР", issue.contract.number if issue.contract else "-"),
                ("Заявка участка", issue.site_request.number if issue.site_request else "-"),
                ("Приходный ордер", issue.stock_receipt.number if issue.stock_receipt else "-"),
                ("Получатель", issue.received_by_name),
                ("Статус", issue.get_status_display()),
            ],
        )
        self._add_table(
            doc,
            ["Код", "Наименование", "Ед.", "Количество", "Цена"],
            [
                [
                    line.material.code,
                    line.material.name,
                    line.material.unit,
                    str(line.quantity),
                    money(line.unit_price),
                ]
                for line in issue.lines.all()
            ],
        )
        self._add_signature(doc, "Кладовщик", "Начальник участка")
        doc.save(path)
        return path

    def _export_writeoff(self, entity_id: int) -> Path:
        act = WriteOffAct.objects.select_related("contract__object").prefetch_related("lines__material").get(pk=entity_id)
        path = self._doc_path("write_off", act.number)
        if self._render_docx_template(DOCX_TEMPLATE_FILES["write_off"], self._writeoff_template_context(act), path):
            return path
        doc = self._prepare_doc("АКТ СПИСАНИЯ МАТЕРИАЛОВ", f"№ {act.number} от {act.act_date}")
        self._add_meta(
            doc,
            [
                ("Договор", act.contract.number),
                ("Объект", act.contract.object.name if act.contract.object else ""),
                ("Участок", act.site_name),
                ("Вид работ", act.work_type),
                ("Объем работ", f"{act.work_volume} {act.volume_unit}".strip()),
                ("Статус", act.get_status_display()),
            ],
        )
        self._add_table(
            doc,
            ["Код", "Наименование", "Ед.", "Норма", "Расчет", "Факт"],
            [
                [
                    line.material.code,
                    line.material.name,
                    line.material.unit,
                    str(line.norm_per_unit),
                    str(line.calculated_quantity),
                    str(line.actual_quantity),
                ]
                for line in act.lines.all()
            ],
        )
        self._add_signature(doc, "Начальник участка", "Начальник монтажного объекта")
        doc.save(path)
        return path

    def _export_work_acceptance(self, entity_id: int) -> Path:
        act = WorkAcceptanceAct.objects.select_related("contract__object").get(pk=entity_id)
        path = self._doc_path("work_acceptance", act.number)
        if self._render_docx_template(DOCX_TEMPLATE_FILES["work_acceptance"], self._work_acceptance_template_context(act), path):
            return path
        customer_name = act.contract.resolved_customer_name() or act.contract.customer_name or "-"
        contractor_name = act.contract.resolved_contractor_name() or self._organization_name() or "-"
        doc = self._prepare_doc("АКТ СДАЧИ-ПРИЕМКИ ВЫПОЛНЕННЫХ РАБОТ", f"№ {act.number} от {act.act_date}")
        self._add_meta(
            doc,
            [
                ("Договор СМР", act.contract.number),
                ("Объект", act.contract.object.name if act.contract.object else act.site_name),
                ("Заказчик", customer_name),
                ("Подрядчик", contractor_name),
                ("Описание работ", act.work_description or act.contract.subject),
                ("Принятый объем", f"{act.accepted_volume or '-'} {act.volume_unit or ''}".strip()),
                ("Сумма", f"{money(act.amount)} руб."),
                ("Статус", act.get_status_display()),
            ],
        )
        self._add_clause(doc, "1", "Работы выполнены в соответствии с договором и переданы Заказчику для приемки.")
        self._add_clause(doc, "2", "Подписание акта подтверждает закрытие выполненного этапа работ по договору СМР.")
        self._add_signature(doc, f"Заказчик: {customer_name}", f"Подрядчик: {contractor_name}")
        doc.save(path)
        return path

    def _export_ppe_issuance(self, entity_id: int) -> Path:
        issuance = PPEIssuance.objects.prefetch_related("lines__worker", "lines__material").get(pk=entity_id)
        path = self._doc_path("ppe_issuance", issuance.number)
        season = (issuance.season or "").casefold()
        if "зим" in season:
            template_name = "Ведомость выдачи спецодежды зимняя_шаблон.docx"
        elif "лет" in season:
            template_name = "Ведомость выдачи спецодежды летняя_шаблон.docx"
        else:
            template_name = "Ведомость выдачи спецодежды перчатки_шаблон.docx"
        if self._render_docx_template(template_name, self._ppe_template_context(issuance), path):
            return path
        doc = self._prepare_doc("ВЕДОМОСТЬ ВЫДАЧИ СПЕЦОДЕЖДЫ", f"№ {issuance.number} от {issuance.issue_date}")
        self._add_meta(
            doc,
            [
                ("Участок", issuance.site_name),
                ("Сезон", issuance.season),
                ("Статус", issuance.get_status_display()),
            ],
        )
        self._add_table(
            doc,
            ["Таб.№", "ФИО", "Материал", "Код", "Размер одежды", "Размер обуви", "Ед.", "Кол-во", "Срок службы, мес."],
            [
                [
                    line.worker.employee_number,
                    line.worker.full_name,
                    line.material.name,
                    line.material.code,
                    line.clothing_size,
                    line.shoe_size,
                    line.material.unit,
                    str(line.quantity),
                    str(line.service_life_months),
                ]
                for line in issuance.lines.all()
            ],
        )
        self._add_signature(doc, "Материально ответственное лицо", "Начальник участка")
        doc.save(path)
        return path

    def _export_supplier_document(self, entity_id: int) -> Path:
        item = SupplierDocument.objects.get(pk=entity_id)
        if item.attachment and Path(item.attachment.path).exists():
            return Path(item.attachment.path)
        path = self._doc_path("supplier_document", item.doc_number)
        template_name = self._supplier_document_template_name(item.doc_type)
        if template_name:
            context = {
                **self._template_common_context(),
                "INVOICE_NUMBER": item.doc_number,
                "INVOICE_DATE": self._date_text(item.doc_date),
                "INVOICE_FACTURE_NUMBER": item.doc_number,
                "INVOICE_FACTURE_DATE": self._date_text(item.doc_date),
                "WAYBILL_NUMBER": item.doc_number,
                "WAYBILL_DATE": self._date_text(item.doc_date),
                "SUPPLIER_NAME": item.supplier.name,
                "SELLER_NAME": item.supplier.name,
                "SHIPPER_NAME": item.supplier.name,
                "SUPPLIER_REQUISITES": self._supplier_requisites(item.supplier),
                "BUYER_NAME": self._organization_name() or "-",
                "CONSIGNEE_NAME": self._organization_name() or "-",
                "PAYER_NAME": self._organization_name() or "-",
                "BUYER_REQUISITES": self._organization_requisites(),
                "TOTAL_AMOUNT": money(item.amount),
                "TOTAL_TO_PAY": money(item.amount),
                "TOTAL_TO_PAY_WORDS": f"{money(item.amount)} руб.",
                "VAT_AMOUNT": money(item.vat_amount),
                "VAT_RATE": "20",
                "BASIS_DOCUMENT": item.request.number if item.request else "",
            }
            if self._render_docx_template(template_name, context, path):
                return path
        doc = self._prepare_doc("ДОКУМЕНТ ПОСТАВКИ", f"{item.doc_type} № {item.doc_number} от {item.doc_date}")
        self._add_meta(
            doc,
            [
                ("Поставщик", item.supplier.name),
                ("Сумма", f"{money(item.amount)} руб."),
                ("НДС", f"{money(item.vat_amount)} руб."),
                ("Комментарий", item.notes or ""),
            ],
        )
        doc.save(path)
        return path
