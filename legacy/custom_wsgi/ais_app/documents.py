from __future__ import annotations

import mimetypes
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from .config import CONFIG
from .database import Database


def money(value: float | int | str | None) -> str:
    numeric = float(value or 0)
    return f"{numeric:,.2f}".replace(",", " ").replace(".", ",")


def fit_columns(sheet) -> None:
    for idx, column_cells in enumerate(sheet.columns, 1):
        max_length = 10
        for cell in column_cells:
            value = "" if cell.value is None else str(cell.value)
            max_length = max(max_length, len(value) + 2)
        sheet.column_dimensions[get_column_letter(idx)].width = min(max_length, 45)


class Exporter:
    def __init__(self, db: Database) -> None:
        self.db = db

    def export_document(self, entity_type: str, entity_id: int) -> Path:
        handlers = {
            "smr_contract": self._export_smr_contract,
            "supply_contract": self._export_supply_contract,
            "procurement_request": self._export_procurement_request,
            "stock_receipt": self._export_stock_receipt,
            "stock_issue": self._export_stock_issue,
            "write_off": self._export_write_off,
            "ppe_issuance": self._export_ppe_issuance,
            "supplier_document": self._export_supplier_document,
        }
        if entity_type not in handlers:
            raise ValueError("Для этого типа документа выгрузка не реализована.")
        path = handlers[entity_type](entity_id)
        self.db.update_document_file(entity_type, entity_id, str(path))
        return path

    def export_report(self, report_name: str, filters: dict[str, str]) -> Path:
        handlers = {
            "stock": (self.db.report_stock, "Остатки материалов"),
            "purchases": (self.db.report_purchases, "Закупки материалов"),
            "writeoffs": (self.db.report_write_offs, "Списание материалов"),
            "work": (self.db.report_work_logs, "Работа участков"),
            "summary": (self.db.report_summary, "Сводный отчет"),
            "ppe": (self.db.report_ppe, "Выданная спецодежда"),
        }
        if report_name not in handlers:
            raise ValueError("Неизвестный отчет.")
        provider, title = handlers[report_name]
        rows = provider(filters)
        date_from = filters.get("date_from") or datetime.now().replace(day=1).date().isoformat()
        date_to = filters.get("date_to") or datetime.now().date().isoformat()
        file_name = f"{report_name}_{date_from}_{date_to}.xlsx".replace(":", "-")
        path = CONFIG.exports_dir / file_name
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Отчет"
        sheet["A1"] = title
        sheet["A1"].font = Font(size=14, bold=True)
        sheet["A2"] = f"Период: {date_from} — {date_to}"
        if rows:
            headers = list(rows[0].keys())
            row_index = 4
            fill = PatternFill("solid", fgColor="D7ECFF")
            for col_index, header in enumerate(headers, 1):
                cell = sheet.cell(row=row_index, column=col_index, value=header)
                cell.font = Font(bold=True)
                cell.fill = fill
                cell.alignment = Alignment(horizontal="center", vertical="center")
            for data_row in rows:
                row_index += 1
                for col_index, header in enumerate(headers, 1):
                    sheet.cell(row=row_index, column=col_index, value=data_row.get(header))
        else:
            sheet["A4"] = "Нет данных за выбранный период."
        fit_columns(sheet)
        workbook.save(path)
        return path

    def content_type(self, path: Path) -> str:
        guessed, _ = mimetypes.guess_type(path.name)
        return guessed or "application/octet-stream"

    def _doc_path(self, prefix: str, number: str) -> Path:
        safe_number = number.replace("/", "_").replace("\\", "_")
        return CONFIG.exports_dir / f"{prefix}_{safe_number}.docx"

    def _prepare_doc(self, title: str, subtitle: str = "") -> Document:
        document = Document()
        style = document.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(14)
        if subtitle:
            subtitle_paragraph = document.add_paragraph()
            subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_paragraph.add_run(subtitle).italic = True
        document.add_paragraph()
        return document

    def _add_meta(self, document: Document, items: Iterable[tuple[str, str]]) -> None:
        for label, value in items:
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{label}: ").bold = True
            paragraph.add_run(str(value))

    def _add_table(self, document: Document, headers: list[str], rows: list[list[str]]) -> None:
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        head_cells = table.rows[0].cells
        for index, header in enumerate(headers):
            head_cells[index].text = header
        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = str(value)

    def _add_signature(self, document: Document, left_label: str, right_label: str) -> None:
        document.add_paragraph()
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.cell(0, 0).text = f"{left_label}\n\n_____________________"
        table.cell(0, 1).text = f"{right_label}\n\n_____________________"

    def _export_smr_contract(self, entity_id: int) -> Path:
        contract = self.db.fetch_one(
            """
            SELECT c.*, o.name AS object_name, o.address
            FROM smr_contracts c
            LEFT JOIN construction_objects o ON o.id = c.object_id
            WHERE c.id = ?
            """,
            (entity_id,),
        )
        if not contract:
            raise ValueError("Договор не найден.")
        doc = self._prepare_doc("ДОГОВОР НА ВЫПОЛНЕНИЕ СМР", f"№ {contract['number']} от {contract['contract_date']}")
        self._add_meta(
            doc,
            [
                ("Заказчик", contract["customer_name"]),
                ("Подрядчик", CONFIG.company_name),
                ("Объект", contract.get("object_name") or ""),
                ("Адрес объекта", contract.get("address") or ""),
                ("Предмет", contract["subject"]),
                ("Вид работ", contract.get("work_type") or ""),
                ("Плановый объем", f"{contract.get('planned_volume') or 0} {contract.get('volume_unit') or ''}".strip()),
                ("Стоимость", f"{money(contract['amount'])} руб."),
                ("Сроки", f"{contract.get('start_date') or '-'} — {contract.get('end_date') or '-'}"),
            ],
        )
        doc.add_paragraph(
            "Подрядчик обязуется выполнить строительные и монтажные работы в полном объеме, а Заказчик — принять и оплатить результат работ в порядке, предусмотренном договором."
        )
        doc.add_paragraph(
            "Используемые материалы учитываются в АИС АО «СТ-1» и привязываются к конкретному объекту строительства, что обеспечивает прозрачность движения ресурсов и последующее списание по нормам."
        )
        self._add_signature(doc, "Заказчик", "Подрядчик")
        path = self._doc_path("smr_contract", contract["number"])
        doc.save(path)
        return path

    def _export_supply_contract(self, entity_id: int) -> Path:
        contract = self.db.fetch_one(
            """
            SELECT c.*, s.name AS supplier_name, sc.number AS smr_number
            FROM supply_contracts c
            JOIN suppliers s ON s.id = c.supplier_id
            LEFT JOIN smr_contracts sc ON sc.id = c.related_smr_contract_id
            WHERE c.id = ?
            """,
            (entity_id,),
        )
        if not contract:
            raise ValueError("Договор поставки не найден.")
        doc = self._prepare_doc("ДОГОВОР ПОСТАВКИ МАТЕРИАЛОВ", f"№ {contract['number']} от {contract['contract_date']}")
        self._add_meta(
            doc,
            [
                ("Поставщик", contract["supplier_name"]),
                ("Покупатель", CONFIG.company_name),
                ("Связанный договор СМР", contract.get("smr_number") or "-"),
                ("Сумма", f"{money(contract['amount'])} руб."),
                ("Статус", contract.get("status") or ""),
            ],
        )
        doc.add_paragraph(contract.get("terms") or "Поставка материалов выполняется по заявкам снабженца через АИС.")
        self._add_signature(doc, "Поставщик", "Покупатель")
        path = self._doc_path("supply_contract", contract["number"])
        doc.save(path)
        return path

    def _export_procurement_request(self, entity_id: int) -> Path:
        request = self.db.fetch_one(
            """
            SELECT r.*, c.number AS contract_number, s.name AS supplier_name
            FROM procurement_requests r
            LEFT JOIN smr_contracts c ON c.id = r.contract_id
            LEFT JOIN suppliers s ON s.id = r.supplier_id
            WHERE r.id = ?
            """,
            (entity_id,),
        )
        if not request:
            raise ValueError("Заявка не найдена.")
        lines = self.db.request_lines(entity_id)
        doc = self._prepare_doc("ЗАЯВКА ПОСТАВЩИКУ НА ЗАКУПКУ МАТЕРИАЛОВ", f"№ {request['number']} от {request['request_date']}")
        self._add_meta(
            doc,
            [
                ("Участок", request["site_name"]),
                ("Договор СМР", request.get("contract_number") or "-"),
                ("Поставщик", request.get("supplier_name") or "-"),
                ("Статус", request["status"]),
            ],
        )
        self._add_table(
            doc,
            ["Код", "Наименование", "Ед.", "Количество", "Цена", "Примечание"],
            [
                [
                    line["material_code"],
                    line["material_name"],
                    line["unit"],
                    str(line["quantity"]),
                    money(line["unit_price"]),
                    line["notes"] or "",
                ]
                for line in lines
            ],
        )
        self._add_signature(doc, "Начальник участка / снабженец", "Начальник монтажного объекта")
        path = self._doc_path("procurement_request", request["number"])
        doc.save(path)
        return path

    def _export_stock_receipt(self, entity_id: int) -> Path:
        receipt = self.db.fetch_one(
            """
            SELECT r.*, s.name AS supplier_name
            FROM stock_receipts r
            JOIN suppliers s ON s.id = r.supplier_id
            WHERE r.id = ?
            """,
            (entity_id,),
        )
        if not receipt:
            raise ValueError("Приходный ордер не найден.")
        lines = self.db.receipt_lines(entity_id)
        doc = self._prepare_doc("ПРИХОДНЫЙ ОРДЕР", f"№ {receipt['number']} от {receipt['receipt_date']}")
        self._add_meta(
            doc,
            [
                ("Поставщик", receipt["supplier_name"]),
                ("Склад", CONFIG.warehouse_name),
                ("Статус", receipt["status"]),
            ],
        )
        self._add_table(
            doc,
            ["Код", "Наименование", "Ед.", "Количество", "Цена", "Сумма"],
            [
                [
                    line["material_code"],
                    line["material_name"],
                    line["unit"],
                    str(line["quantity"]),
                    money(line["unit_price"]),
                    money(float(line["quantity"]) * float(line["unit_price"])),
                ]
                for line in lines
            ],
        )
        self._add_signature(doc, "Кладовщик", "Материально-ответственное лицо")
        path = self._doc_path("stock_receipt", receipt["number"])
        doc.save(path)
        return path

    def _export_stock_issue(self, entity_id: int) -> Path:
        issue = self.db.fetch_one(
            """
            SELECT i.*, c.number AS contract_number
            FROM stock_issues i
            LEFT JOIN smr_contracts c ON c.id = i.contract_id
            WHERE i.id = ?
            """,
            (entity_id,),
        )
        if not issue:
            raise ValueError("Требование-накладная не найдена.")
        lines = self.db.issue_lines(entity_id)
        doc = self._prepare_doc("ТРЕБОВАНИЕ-НАКЛАДНАЯ", f"№ {issue['number']} от {issue['issue_date']}")
        self._add_meta(
            doc,
            [
                ("Участок", issue["site_name"]),
                ("Договор СМР", issue.get("contract_number") or "-"),
                ("Получатель", issue["received_by_name"]),
                ("Статус", issue["status"]),
            ],
        )
        self._add_table(
            doc,
            ["Код", "Наименование", "Ед.", "Количество", "Цена"],
            [[line["material_code"], line["material_name"], line["unit"], str(line["quantity"]), money(line["unit_price"])] for line in lines],
        )
        self._add_signature(doc, "Кладовщик", "Начальник участка")
        path = self._doc_path("stock_issue", issue["number"])
        doc.save(path)
        return path

    def _export_write_off(self, entity_id: int) -> Path:
        act = self.db.fetch_one(
            """
            SELECT a.*, c.number AS contract_number, o.name AS object_name
            FROM write_off_acts a
            JOIN smr_contracts c ON c.id = a.contract_id
            LEFT JOIN construction_objects o ON o.id = c.object_id
            WHERE a.id = ?
            """,
            (entity_id,),
        )
        if not act:
            raise ValueError("Акт списания не найден.")
        lines = self.db.write_off_lines(entity_id)
        doc = self._prepare_doc("АКТ СПИСАНИЯ МАТЕРИАЛОВ", f"№ {act['number']} от {act['act_date']}")
        self._add_meta(
            doc,
            [
                ("Договор", act["contract_number"]),
                ("Объект", act.get("object_name") or ""),
                ("Участок", act["site_name"]),
                ("Вид работ", act["work_type"]),
                ("Объем работ", f"{act['work_volume']} {act.get('volume_unit') or ''}".strip()),
                ("Статус", act["status"]),
            ],
        )
        self._add_table(
            doc,
            ["Код", "Наименование", "Ед.", "Норма", "Расчет", "Факт"],
            [
                [
                    line["material_code"],
                    line["material_name"],
                    line["unit"],
                    str(line["norm_per_unit"]),
                    str(line["calculated_quantity"]),
                    str(line["actual_quantity"]),
                ]
                for line in lines
            ],
        )
        doc.add_paragraph("Списание рассчитано автоматически на основании норм расхода и объема выполненных работ.")
        self._add_signature(doc, "Начальник участка", "Начальник монтажного объекта")
        path = self._doc_path("write_off", act["number"])
        doc.save(path)
        return path

    def _export_ppe_issuance(self, entity_id: int) -> Path:
        issuance = self.db.fetch_one("SELECT * FROM ppe_issuances WHERE id = ?", (entity_id,))
        if not issuance:
            raise ValueError("Ведомость выдачи не найдена.")
        lines = self.db.ppe_lines(entity_id)
        doc = self._prepare_doc("ВЕДОМОСТЬ УЧЕТА ВЫДАЧИ СПЕЦОДЕЖДЫ", f"№ {issuance['number']} от {issuance['issue_date']}")
        self._add_meta(
            doc,
            [
                ("Участок", issuance["site_name"]),
                ("Сезон", issuance.get("season") or ""),
                ("Статус", issuance["status"]),
            ],
        )
        self._add_table(
            doc,
            ["Таб.№", "ФИО", "Материал", "Код", "Ед.", "Кол-во", "Срок службы, мес."],
            [
                [
                    line["employee_number"],
                    line["full_name"],
                    line["material_name"],
                    line["material_code"],
                    line["unit"],
                    str(line["quantity"]),
                    str(line["service_life_months"]),
                ]
                for line in lines
            ],
        )
        self._add_signature(doc, "Материально ответственное лицо", "Начальник участка")
        path = self._doc_path("ppe_issuance", issuance["number"])
        doc.save(path)
        return path

    def _export_supplier_document(self, entity_id: int) -> Path:
        item = self.db.fetch_one("SELECT * FROM supplier_documents WHERE id = ?", (entity_id,))
        if not item:
            raise ValueError("Документ поставщика не найден.")
        if item.get("stored_path") and Path(item["stored_path"]).exists():
            return Path(item["stored_path"])
        doc = self._prepare_doc("ДОКУМЕНТ ПОСТАВКИ", f"{item['doc_type']} № {item['doc_number']} от {item['doc_date']}")
        self._add_meta(
            doc,
            [
                ("Сумма", f"{money(item['amount'])} руб."),
                ("НДС", f"{money(item['vat_amount'])} руб."),
                ("Примечание", item.get("notes") or ""),
            ],
        )
        path = self._doc_path("supplier_document", item["doc_number"])
        doc.save(path)
        return path
